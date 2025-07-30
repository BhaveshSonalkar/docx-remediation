#!/usr/bin/env python3
"""
Hardcoded DOCX snippet generators for each accessibility issue
"""

from docx import Document
import docx.shared
import base64
from io import BytesIO

def docx_to_base64(doc):
    """Helper function to convert DOCX document to base64"""
    try:
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        docx_data = buffer.getvalue()
        return base64.b64encode(docx_data).decode('utf-8')
    except Exception as e:
        print(f"Error converting DOCX to base64: {e}")
        return None

# Issue 1: Title contrast
def create_title_original_snippet():
    """Title with poor contrast (original problem)"""
    doc = Document()
    title = doc.add_heading('Sample Document with Accessibility Issues', 0)
    if title.runs:
        title.runs[0].font.color.rgb = docx.shared.RGBColor(200, 200, 200)  # Poor contrast gray
    return docx_to_base64(doc)

def create_title_fixed_snippet():
    """Title with good contrast (fixed)"""
    doc = Document()
    title = doc.add_heading('Sample Document with Accessibility Issues (High Contrast Version)', 0)
    if title.runs:
        title.runs[0].font.color.rgb = docx.shared.RGBColor(0, 0, 0)  # High contrast black
        title.runs[0].font.bold = True
    return docx_to_base64(doc)

# Issue 2: Paragraph should be heading
def create_paragraph_original_snippet():
    """Paragraph that should be heading (original problem)"""
    doc = Document()
    doc.add_paragraph('This is a paragraph that should be a heading.')  # Wrong: paragraph
    return docx_to_base64(doc)

def create_paragraph_fixed_snippet():
    """Paragraph converted to proper heading (fixed)"""
    doc = Document()
    heading = doc.add_heading('Introduction', level=2)  # Fixed: proper H2
    if heading.runs:
        heading.runs[0].font.bold = True
    return docx_to_base64(doc)

# Issue 3: Wrong heading hierarchy
def create_hierarchy_original_snippet():
    """Wrong heading hierarchy (original problem)"""
    doc = Document()
    doc.add_heading('Subsection', level=3)  # Wrong: H3 without H1/H2
    return docx_to_base64(doc)

def create_hierarchy_fixed_snippet():
    """Proper heading hierarchy (fixed)"""
    doc = Document()
    heading = doc.add_heading('Key Features and Benefits', level=2)  # Fixed: proper H2
    if heading.runs:
        heading.runs[0].font.bold = True
    return docx_to_base64(doc)

# Issue 4: Text contrast
def create_text_original_snippet():
    """Text with poor contrast (original problem)"""
    doc = Document()
    p = doc.add_paragraph('This text has insufficient color contrast.')
    if p.runs:
        p.runs[0].font.color.rgb = docx.shared.RGBColor(180, 180, 180)  # Very light gray
    return docx_to_base64(doc)

def create_text_fixed_snippet():
    """Text with good contrast (fixed)"""
    doc = Document()
    p = doc.add_paragraph('This text has been optimized for excellent color contrast and readability.')
    if p.runs:
        p.runs[0].font.color.rgb = docx.shared.RGBColor(0, 0, 0)  # High contrast black
        p.runs[0].font.size = docx.shared.Pt(12)
        p.runs[0].font.name = 'Calibri'
    return docx_to_base64(doc)

# Issue 5: Alt text
def create_alt_original_snippet():
    """Missing alt text reference (original problem)"""
    doc = Document()
    doc.add_paragraph('Please refer to the chart below for more information.')  # Vague reference
    return docx_to_base64(doc)

def create_alt_fixed_snippet():
    """Descriptive content (fixed)"""
    doc = Document()
    p = doc.add_paragraph('Please refer to the Annual Sales Performance Chart below, which shows quarterly revenue trends for 2023.')
    if p.runs:
        p.runs[0].font.size = docx.shared.Pt(11)
    return docx_to_base64(doc)

# Issue 6: Table headers
def create_table_original_snippet():
    """Table without headers (original problem)"""
    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f'Data {i+1}-{j+1}'  # No headers, just data
    return docx_to_base64(doc)

def create_table_fixed_snippet():
    """Table with proper headers (fixed)"""
    doc = Document()
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    
    # Add headers with bold formatting
    headers = ['Product Name', 'Price', 'Availability']
    header_row = table.rows[0]
    for col_idx, header_text in enumerate(headers):
        cell = header_row.cells[col_idx]
        cell.text = header_text
        # Make header bold
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add sample data rows
    sample_data = [
        ['Product A', '$19.99', 'Available'],
        ['Product B', '$29.99', 'Out of Stock']
    ]
    for row_idx, row_data in enumerate(sample_data, 1):
        for col_idx, data in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = data
    
    return docx_to_base64(doc)

# Issue 7: Link text
def create_link_original_snippet():
    """Non-descriptive link text (original problem)"""
    doc = Document()
    p = doc.add_paragraph('Click ')
    link_run = p.add_run('here')  # Problem: vague link text
    link_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    link_run.font.underline = True
    p.add_run(' for more information.')
    return docx_to_base64(doc)

def create_link_fixed_snippet():
    """Descriptive link text (fixed)"""
    doc = Document()
    p = doc.add_paragraph('Click ')
    link_run = p.add_run('download the accessibility report')  # Fixed: descriptive text
    link_run.font.color.rgb = docx.shared.RGBColor(0, 0, 255)
    link_run.font.underline = True
    p.add_run(' for more information.')
    return docx_to_base64(doc)

# Issue 8: Font size
def create_font_original_snippet():
    """Text too small (original problem)"""
    doc = Document()
    p = doc.add_paragraph('This text is too small to read easily.')
    if p.runs:
        p.runs[0].font.size = docx.shared.Pt(6)  # Very small font
    return docx_to_base64(doc)

def create_font_fixed_snippet():
    """Readable font size (fixed)"""
    doc = Document()
    p = doc.add_paragraph('This text is now sized appropriately for easy reading and accessibility compliance.')
    if p.runs:
        p.runs[0].font.size = docx.shared.Pt(12)  # Readable size
        p.runs[0].font.name = 'Calibri'
    return docx_to_base64(doc)

def get_hardcoded_snippets(issue_id, issue_context):
    """Get hardcoded DOCX snippets based on issue type"""
    try:
        clause = issue_context.get('clause', '')
        element_xpath = issue_context.get('element_xpath', '')
        description = issue_context.get('description', '').lower()  # Make case-insensitive
        
        # Map issues to snippet generators based on clause and xpath
        if '//w:p[1]' in element_xpath and 'color contrast' in description:
            # Issue 1: Title contrast
            return {
                'original': create_title_original_snippet(),
                'fixed': create_title_fixed_snippet()
            }
        elif '//w:p[2]' in element_xpath and 'paragraph' in description and 'heading' in description:
            # Issue 2: Paragraph to heading
            return {
                'original': create_paragraph_original_snippet(),
                'fixed': create_paragraph_fixed_snippet()
            }
        elif '//w:p[3]' in element_xpath and 'heading hierarchy' in description:
            # Issue 3: Heading hierarchy
            return {
                'original': create_hierarchy_original_snippet(),
                'fixed': create_hierarchy_fixed_snippet()
            }
        elif '//w:p[4]' in element_xpath and 'color contrast' in description:
            # Issue 4: Text contrast
            return {
                'original': create_text_original_snippet(),
                'fixed': create_text_fixed_snippet()
            }
        elif '//w:p[5]' in element_xpath and ('alt' in description or 'alternative' in description):
            # Issue 5: Alt text
            return {
                'original': create_alt_original_snippet(),
                'fixed': create_alt_fixed_snippet()
            }
        elif '//w:tbl[1]' in element_xpath and 'table' in description:
            # Issue 6: Table headers
            return {
                'original': create_table_original_snippet(),
                'fixed': create_table_fixed_snippet()
            }
        elif '//w:p[6]' in element_xpath and 'link' in description:
            # Issue 7: Link text
            return {
                'original': create_link_original_snippet(),
                'fixed': create_link_fixed_snippet()
            }
        elif '//w:p[7]' in element_xpath and 'too small' in description:
            # Issue 8: Font size
            return {
                'original': create_font_original_snippet(),
                'fixed': create_font_fixed_snippet()
            }
        else:
            # Default fallback
            doc = Document()
            doc.add_paragraph('Default snippet - issue type not recognized')
            default_snippet = docx_to_base64(doc)
            return {
                'original': default_snippet,
                'fixed': default_snippet
            }
            
    except Exception as e:
        print(f"Error getting hardcoded snippets: {e}")
        return None