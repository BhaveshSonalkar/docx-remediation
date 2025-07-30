from docx import Document
from docx.shared import RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_test_document():
    """Create a test DOCX file with various accessibility issues"""
    doc = Document()
    
    # Issue 1: Title with poor contrast
    title = doc.add_heading('Sample Document with Accessibility Issues', 0)
    title_run = title.runs[0]
    title_run.font.color.rgb = RGBColor(200, 200, 200)  # Light gray - poor contrast
    
    # Issue 2: Missing heading structure
    doc.add_paragraph('This is a paragraph that should be a heading.')
    
    # Issue 3: Heading 3 without proper hierarchy
    doc.add_heading('Subsection', level=3)
    
    # Issue 4: Text with insufficient contrast
    p1 = doc.add_paragraph('This text has insufficient color contrast. ')
    run1 = p1.runs[0]
    run1.font.color.rgb = RGBColor(180, 180, 180)  # Very light gray
    
    # Issue 5: Missing alt text for image reference
    p2 = doc.add_paragraph('Please refer to the chart below for more information.')
    
    # Issue 6: Table without headers
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell.text = f'Data {i+1}-{j+1}'
    
    # Issue 7: Link without descriptive text
    p3 = doc.add_paragraph('Click ')
    link_run = p3.add_run('here')
    link_run.font.color.rgb = RGBColor(0, 0, 255)
    p3.add_run(' for more information.')
    
    # Issue 8: Text that's too small
    p4 = doc.add_paragraph('This text is too small to read easily.')
    small_run = p4.runs[0]
    small_run.font.size = Inches(0.08)  # Very small font
    
    # Save the document
    doc.save('test_document_with_issues.docx')
    print("Test document created: test_document_with_issues.docx")

if __name__ == '__main__':
    create_test_document()