from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import uuid
import shutil
from datetime import datetime
from werkzeug.utils import secure_filename
from docx import Document
from hardcoded_snippets import get_hardcoded_snippets

app = Flask(__name__)
CORS(app)

# Configuration
UPLOAD_FOLDER = os.path.abspath(os.path.join(os.path.dirname(__file__), '.', 'uploads'))
ALLOWED_EXTENSIONS = {'docx'}

# Ensure upload directory exists
os.makedirs(UPLOAD_FOLDER, exist_ok=True)

# In-memory storage for prototype
documents = {}
accessibility_issues = {}
staged_changes = {}  # Add this for storing staged changes

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def create_backup(file_path):
    """Create a backup of the original document"""
    backup_path = file_path.replace('.docx', '_backup.docx')
    if not os.path.exists(backup_path):
        shutil.copy2(file_path, backup_path)
    return backup_path

def parse_xpath_index(xpath_part):
    """Extract index from XPath part like 'w:p[2]' -> 2"""
    if '[' in xpath_part and ']' in xpath_part:
        index_str = xpath_part.split('[')[1].split(']')[0]
        return int(index_str) - 1  # Convert to 0-based index
    return 0

def find_element_by_xpath(doc, element_xpath):
    """Find document element using XPath-like selector"""
    if not element_xpath:
        return None
    
    try:
        # Parse common DOCX XPath patterns
        # Examples: //w:p[1], //w:p[2]/w:r[1], //w:tbl[1], //w:tbl[1]/w:tr[1]/w:tc[2]
        
        if '//w:p[' in element_xpath:
            # Paragraph element
            paragraph_match = element_xpath.split('//w:p[')[1].split(']')[0]
            paragraph_index = int(paragraph_match) - 1  # Convert to 0-based
            
            if paragraph_index < len(doc.paragraphs):
                paragraph = doc.paragraphs[paragraph_index]
                
                # Check if it's targeting a specific run within the paragraph
                if '/w:r[' in element_xpath:
                    run_match = element_xpath.split('/w:r[')[1].split(']')[0]
                    run_index = int(run_match) - 1
                    if run_index < len(paragraph.runs):
                        return paragraph.runs[run_index]
                
                return paragraph
        
        elif '//w:tbl[' in element_xpath:
            # Table element
            table_match = element_xpath.split('//w:tbl[')[1].split(']')[0]
            table_index = int(table_match) - 1  # Convert to 0-based
            
            if table_index < len(doc.tables):
                table = doc.tables[table_index]
                
                # Check if it's targeting a specific row/cell
                if '/w:tr[' in element_xpath:
                    row_match = element_xpath.split('/w:tr[')[1].split(']')[0]
                    row_index = int(row_match) - 1
                    
                    if row_index < len(table.rows):
                        row = table.rows[row_index]
                        
                        if '/w:tc[' in element_xpath:
                            cell_match = element_xpath.split('/w:tc[')[1].split(']')[0]
                            cell_index = int(cell_match) - 1
                            
                            if cell_index < len(row.cells):
                                return row.cells[cell_index]
                        
                        return row
                
                return table
        
        return None
        
    except (ValueError, IndexError) as e:
        print(f"Error parsing XPath '{element_xpath}': {e}")
        return None

def find_element_by_content(doc, target_content, element_type='paragraph'):
    """Find document element by content text (fallback method)"""
    target_content = target_content.strip()
    
    if element_type == 'paragraph':
        for paragraph in doc.paragraphs:
            if paragraph.text.strip() == target_content:
                return paragraph
            # Also check for partial matches in case of formatting
            if target_content in paragraph.text:
                return paragraph
    
    elif element_type == 'heading':
        for paragraph in doc.paragraphs:
            if paragraph.style.name.startswith('Heading') and paragraph.text.strip() == target_content:
                return paragraph
    
    elif element_type == 'table_cell':
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip() == target_content:
                        return cell
    
    return None

def modify_document_element(doc, original_content, new_content, element_xpath=None):
    """Modify a document element with new content using XPath when available"""
    try:
        element = None
        
        # First, try to find element using XPath if provided
        if element_xpath:
            element = find_element_by_xpath(doc, element_xpath)
            print(f"XPath '{element_xpath}' -> Found element: {type(element).__name__ if element else 'None'}")
        
        # If XPath didn't work, fall back to content-based matching
        if not element:
            print(f"Falling back to content-based search for: '{original_content[:50]}...'")
            
            # Try to find the element by content
            element = find_element_by_content(doc, original_content, 'paragraph')
            
            if not element:
                # Try finding in headings
                element = find_element_by_content(doc, original_content, 'heading')
            
            if not element:
                # Try finding in table cells
                element = find_element_by_content(doc, original_content, 'table_cell')
        
        # Modify the element if found
        if element:
            element_type = type(element).__name__
            print(f"Modifying {element_type}: '{original_content[:30]}...' -> '{new_content[:30]}...'")
            
            if hasattr(element, 'clear') and hasattr(element, 'add_run'):
                # Paragraph or Run element
                element.clear()
                element.add_run(new_content)
                return True
            elif hasattr(element, 'text'):
                # Table cell or other text element
                element.text = new_content
                return True
            elif hasattr(element, '_element'):
                # Handle run elements specifically
                if element_type == 'Run':
                    element.text = new_content
                    return True
                # For paragraph-like elements with runs
                elif hasattr(element, 'runs') and element.runs:
                    # Clear existing runs and add new content
                    for run in element.runs:
                        run.clear()
                    element.runs[0].text = new_content if element.runs else element.add_run(new_content)
                    return True
            
            print(f"Warning: Don't know how to modify element type: {element_type}")
            return False
        
        print(f"Element not found for content: '{original_content[:50]}...'")
        return False
        
    except Exception as e:
        print(f"Error modifying element: {e}")
        import traceback
        traceback.print_exc()
        return False

def apply_changes_to_docx(file_path, changes):
    """Apply multiple changes to a DOCX document"""
    try:
        # Create backup first
        backup_path = create_backup(file_path)
        
        # Load the document
        doc = Document(file_path)
        
        applied_changes = []
        failed_changes = []
        
        for change in changes:
            original_content = change['original_content']
            new_content = change['new_content']
            element_xpath = change.get('element_xpath', '')
            
            success = modify_document_element(doc, original_content, new_content, element_xpath)
            
            if success:
                applied_changes.append(change['id'])
            else:
                failed_changes.append({
                    'change_id': change['id'],
                    'reason': 'Element not found or could not be modified'
                })
        
        # Save the modified document
        doc.save(file_path)
        
        return {
            'success': True,
            'applied_changes': applied_changes,
            'failed_changes': failed_changes,
            'backup_path': backup_path,
            'total_applied': len(applied_changes),
            'total_failed': len(failed_changes)
        }
        
    except Exception as e:
        return {
            'success': False,
            'error': str(e),
            'applied_changes': [],
            'failed_changes': [{'change_id': change['id'], 'reason': str(e)} for change in changes]
        }

@app.route('/api/documents/upload', methods=['POST'])
def upload_document():
    """Upload DOCX file and store metadata"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file provided'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No file selected'}), 400
    
    if not allowed_file(file.filename):
        return jsonify({'error': 'Only DOCX files are allowed'}), 400
    
    try:
        # Generate unique document ID
        document_id = str(uuid.uuid4())
        
        # Save file with unique name
        filename = secure_filename(file.filename)
        file_path = os.path.join(UPLOAD_FOLDER, f"{document_id}_{filename}")
        file.save(file_path)
        
        # Store document metadata
        documents[document_id] = {
            'id': document_id,
            'filename': filename,
            'file_path': file_path,
            'upload_date': datetime.now().isoformat(),
            'status': 'uploaded'
        }
        
        return jsonify({
            'document_id': document_id,
            'filename': filename,
            'status': 'uploaded'
        }), 201
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/documents/<document_id>/file', methods=['GET'])
def get_document_file(document_id):
    """Return the original DOCX file for client-side rendering"""
    if document_id not in documents:
        return jsonify({'error': 'Document not found'}), 404
    
    try:
        document = documents[document_id]
        file_path = document['file_path']
        
        # Update document status to ready when file is accessed
        documents[document_id]['status'] = 'ready'
        
        return send_file(
            file_path,
            as_attachment=False,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/documents/<document_id>/scan', methods=['GET'])
def scan_document(document_id):
    """Trigger accessibility scan (hardcoded for prototype)"""
    if document_id not in documents:
        return jsonify({'error': 'Document not found'}), 404
    
    try:
        # Update document status
        documents[document_id]['status'] = 'scanning'
        
        # Hardcoded accessibility issues for prototype
        # In real implementation, this would call actual accessibility scanning service
        # Comprehensive accessibility issues for prototype testing
        hardcoded_issues = [
            {
                'id': str(uuid.uuid4()),
                'document_id': document_id,
                'clause': 'WCAG 2.1 AA 1.4.3',
                'description': 'Insufficient color contrast in document title',
                'status': 'active',
                'wcag_level': 'AA',
                'details': {
                    'contrast_ratio': 1.53,
                    'required_ratio': 4.5,
                    'foreground_color': '#C8C8C8',
                    'background_color': '#FFFFFF',
                    'original_content': 'Sample Document with Accessibility Issues',
                    'element_xpath': '//w:p[1]/w:r[1]/w:t'
                },
                'element_xpath': '//w:p[1]/w:r[1]',
                'is_fixed': False
            },
            {
                'id': str(uuid.uuid4()),
                'document_id': document_id,
                'clause': 'WCAG 2.1 A 1.3.1',
                'description': 'Missing heading structure - paragraph should be a heading',
                'status': 'active',
                'wcag_level': 'A',
                'details': {
                    'issue_type': 'heading_structure',
                    'found_element': 'paragraph',
                    'expected_element': 'heading',
                    'original_content': 'This is a paragraph that should be a heading.',
                    'element_xpath': '//w:p[2]/w:r[1]/w:t'
                },
                'element_xpath': '//w:p[2]',
                'is_fixed': False
            },
            {
                'id': str(uuid.uuid4()),
                'document_id': document_id,
                'clause': 'WCAG 2.1 A 1.3.1',
                'description': 'Improper heading hierarchy - h3 without preceding h2',
                'status': 'active',
                'wcag_level': 'A',
                'details': {
                    'issue_type': 'heading_hierarchy',
                    'found_level': 'h3',
                    'expected_level': 'h2',
                    'original_content': 'Subsection',
                    'element_xpath': '//w:p[3]/w:r[1]/w:t'
                },
                'element_xpath': '//w:p[3]',
                'is_fixed': False
            },
            {
                'id': str(uuid.uuid4()),
                'document_id': document_id,
                'clause': 'WCAG 2.1 AA 1.4.3',
                'description': 'Insufficient color contrast in body text',
                'status': 'active',
                'wcag_level': 'AA',
                'details': {
                    'contrast_ratio': 1.23,
                    'required_ratio': 4.5,
                    'foreground_color': '#B4B4B4',
                    'background_color': '#FFFFFF',
                    'original_content': 'This text has insufficient color contrast.',
                    'element_xpath': '//w:p[4]/w:r[1]/w:t'
                },
                'element_xpath': '//w:p[4]/w:r[1]',
                'is_fixed': False
            },
            {
                'id': str(uuid.uuid4()),
                'document_id': document_id,
                'clause': 'WCAG 2.1 A 1.1.1',
                'description': 'Missing alternative text for referenced image',
                'status': 'active',
                'wcag_level': 'A',
                'details': {
                    'issue_type': 'missing_alt_text',
                    'reference_text': 'chart below',
                    'original_content': 'Please refer to the chart below for more information.',
                    'element_xpath': '//w:p[5]/w:r[1]/w:t'
                },
                'element_xpath': '//w:p[5]',
                'is_fixed': False
            },
            {
                'id': str(uuid.uuid4()),
                'document_id': document_id,
                'clause': 'WCAG 2.1 A 1.3.1',
                'description': 'Table missing header row',
                'status': 'active',
                'wcag_level': 'A',
                'details': {
                    'issue_type': 'table_headers',
                    'table_rows': 3,
                    'table_columns': 3,
                    'original_content': 'Data table without headers',
                    'element_xpath': '//w:tbl[1]'
                },
                'element_xpath': '//w:tbl[1]',
                'is_fixed': False
            },
            {
                'id': str(uuid.uuid4()),
                'document_id': document_id,
                'clause': 'WCAG 2.1 A 2.4.4',
                'description': 'Link text not descriptive - "here" is not meaningful',
                'status': 'active',
                'wcag_level': 'A',
                'details': {
                    'issue_type': 'link_text',
                    'link_text': 'here',
                    'context': 'Click here for more information.',
                    'original_content': 'here',
                    'element_xpath': '//w:p[6]/w:r[2]/w:t'
                },
                'element_xpath': '//w:p[6]/w:r[2]',
                'is_fixed': False
            },
            {
                'id': str(uuid.uuid4()),
                'document_id': document_id,
                'clause': 'WCAG 2.1 AA 1.4.4',
                'description': 'Text too small to read without zooming',
                'status': 'active',
                'wcag_level': 'AA',
                'details': {
                    'issue_type': 'font_size',
                    'current_size': '6pt',
                    'minimum_size': '12pt',
                    'original_content': 'This text is too small to read easily.',
                    'element_xpath': '//w:p[7]/w:r[1]/w:t'
                },
                'element_xpath': '//w:p[7]/w:r[1]',
                'is_fixed': False
            }
        ]        
        # Store issues in memory
        for issue in hardcoded_issues:
            accessibility_issues[issue['id']] = issue
        
        # Update document status
        documents[document_id]['status'] = 'ready'
        
        return jsonify({
            'scan_results': hardcoded_issues,
            'document_id': document_id,
            'total_issues': len(hardcoded_issues)
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/issues/<document_id>', methods=['GET'])
def get_document_issues(document_id):
    """Get all accessibility issues for a document"""
    if document_id not in documents:
        return jsonify({'error': 'Document not found'}), 404
    
    # Filter issues for this document
    document_issues = [
        issue for issue in accessibility_issues.values() 
        if issue['document_id'] == document_id
    ]
    
    return jsonify(document_issues)

@app.route('/api/documents', methods=['GET'])
def list_documents():
    """List all uploaded documents"""
    return jsonify(list(documents.values()))

@app.route('/api/health', methods=['GET'])
def health_check():
    """Health check endpoint"""
    return jsonify({'status': 'healthy', 'service': 'docx-remediation-backend'})

@app.route('/api/issues/<issue_id>/suggest-fix', methods=['POST'])
def suggest_fix(issue_id):
    """Get AI-suggested fix for an issue (hardcoded for prototype)"""
    if issue_id not in accessibility_issues:
        return jsonify({'error': 'Issue not found'}), 404
    
    issue = accessibility_issues[issue_id]
    
    # Get the issue details to provide contextual suggestions
    issue_details = issue.get('details', {})
    original_content = issue_details.get('original_content', '')
    clause = issue['clause']
    
    # Create suggestions based on issue clause and original content
    suggestions = {}
    
    if clause == 'WCAG 2.1 AA 1.4.3' and 'Sample Document' in original_content:
        # Title contrast issue
        suggestions = {
            'suggested_text': 'Improve title readability by ensuring sufficient color contrast',
            'confidence': 0.95,
            'fix_type': 'content_improvement',
            'old_value': original_content,
            'new_value': 'Sample Document with Accessibility Issues (High Contrast Version)',
            'element_xpath': issue.get('element_xpath', '')
        }
    elif clause == 'WCAG 2.1 A 1.3.1' and 'paragraph that should be a heading' in original_content:
        # Heading structure issue
        suggestions = {
            'suggested_text': 'Convert this paragraph to a proper heading for better document structure',
            'confidence': 0.92,
            'fix_type': 'heading_conversion',
            'old_value': original_content,
            'new_value': 'Introduction',
            'element_xpath': issue.get('element_xpath', '')
        }
    elif clause == 'WCAG 2.1 A 1.3.1' and original_content == 'Subsection':
        # Heading hierarchy issue
        suggestions = {
            'suggested_text': 'Improve heading hierarchy by using a more descriptive heading',
            'confidence': 0.88,
            'fix_type': 'heading_improvement',
            'old_value': original_content,
            'new_value': 'Key Features and Benefits',
            'element_xpath': issue.get('element_xpath', '')
        }
    elif clause == 'WCAG 2.1 AA 1.4.3' and 'insufficient color contrast' in original_content:
        # Body text contrast issue
        suggestions = {
            'suggested_text': 'Rewrite text with better contrast-friendly language and clearer messaging',
            'confidence': 0.94,
            'fix_type': 'content_clarity',
            'old_value': original_content,
            'new_value': 'This text has been optimized for excellent color contrast and readability.',
            'element_xpath': issue.get('element_xpath', '')
        }
    elif clause == 'WCAG 2.1 A 1.1.1' and 'chart below' in original_content:
        # Missing alt text issue
        suggestions = {
            'suggested_text': 'Add descriptive text that explains what the chart contains',
            'confidence': 0.89,
            'fix_type': 'descriptive_content',
            'old_value': original_content,
            'new_value': 'Please refer to the Annual Sales Performance Chart below, which shows quarterly revenue trends for 2023.',
            'element_xpath': issue.get('element_xpath', '')
        }
    elif clause == 'WCAG 2.1 A 1.3.1' and 'Data table' in original_content:
        # Table headers issue
        suggestions = {
            'suggested_text': 'Replace generic table description with proper header content',
            'confidence': 0.91,
            'fix_type': 'table_structure',
            'old_value': original_content,
            'new_value': 'Product Name | Price | Availability',
            'element_xpath': issue.get('element_xpath', '')
        }
    elif clause == 'WCAG 2.1 A 2.4.4' and original_content == 'here':
        # Link text issue
        suggestions = {
            'suggested_text': 'Replace vague link text with descriptive text',
            'confidence': 0.96,
            'fix_type': 'link_improvement',
            'old_value': original_content,
            'new_value': 'download the accessibility report',
            'element_xpath': issue.get('element_xpath', '')
        }
    elif clause == 'WCAG 2.1 AA 1.4.4' and 'too small to read' in original_content:
        # Font size issue
        suggestions = {
            'suggested_text': 'Rewrite with emphasis on readability and clear communication',
            'confidence': 0.87,
            'fix_type': 'readability_improvement',
            'old_value': original_content,
            'new_value': 'This text is now sized appropriately for easy reading and accessibility compliance.',
            'element_xpath': issue.get('element_xpath', '')
        }
    else:
        # Default fallback suggestion
        suggestions = {
            'suggested_text': 'Manual review and correction recommended for this accessibility issue',
            'confidence': 0.5,
            'fix_type': 'manual_review',
            'old_value': original_content,
            'new_value': original_content + ' (Please review and correct manually)',
            'element_xpath': issue.get('element_xpath', '')
        }
    
    # If no specific suggestion was created, use the suggestions object directly
    # (since we already created it based on the clause and content)
    suggestion = suggestions
    
    # Generate hardcoded DOCX snippets
    hardcoded_snippets = get_hardcoded_snippets(issue_id, issue)
    original_snippet = hardcoded_snippets['original'] if hardcoded_snippets else None
    fixed_snippet = hardcoded_snippets['fixed'] if hardcoded_snippets else None
    
    response_data = {
        'issue_id': issue_id,
        'suggested_text': suggestion['suggested_text'],
        'confidence': suggestion['confidence'],
        'fix_type': suggestion['fix_type'],
        'old_value': suggestion['old_value'],
        'new_value': suggestion['new_value'],
        'element_xpath': suggestion['element_xpath']
    }
    
    # Add DOCX snippets if successfully generated
    if original_snippet and fixed_snippet:
        response_data['docx_snippets'] = {
            'original': original_snippet,
            'fixed': fixed_snippet
        }
    
    return jsonify(response_data)

def calculate_diff(original, new_content):
    """Calculate detailed diff between original and new content"""
    if not original and not new_content:
        return {'type': 'no_change', 'changes': []}
    
    # Simple line-by-line comparison for prototype
    original_lines = original.split('\n') if original else ['']
    new_lines = new_content.split('\n') if new_content else ['']
    
    changes = []
    max_lines = max(len(original_lines), len(new_lines))
    
    for i in range(max_lines):
        original_line = original_lines[i] if i < len(original_lines) else ''
        new_line = new_lines[i] if i < len(new_lines) else ''
        
        if original_line != new_line:
            change = {
                'line_number': i + 1,
                'type': 'modified' if original_line and new_line else ('added' if new_line else 'deleted'),
                'original': original_line,
                'new': new_line
            }
            changes.append(change)
    
    return {
        'type': 'text_change',
        'changes': changes,
        'summary': {
            'total_changes': len(changes),
            'added_lines': len([c for c in changes if c['type'] == 'added']),
            'deleted_lines': len([c for c in changes if c['type'] == 'deleted']),
            'modified_lines': len([c for c in changes if c['type'] == 'modified'])
        },
        'preview': {
            'original': original[:100] + '...' if len(original) > 100 else original,
            'new': new_content[:100] + '...' if len(new_content) > 100 else new_content
        }
    }

@app.route('/api/issues/<issue_id>/stage-change', methods=['POST'])
def stage_change(issue_id):
    """Stage a fix for an issue"""
    if issue_id not in accessibility_issues:
        return jsonify({'error': 'Issue not found'}), 404
    
    data = request.get_json()
    if not data or 'new_content' not in data:
        return jsonify({'error': 'new_content is required'}), 400
    
    # Validate content
    new_content = data['new_content'].strip()
    if not new_content:
        return jsonify({'error': 'new_content cannot be empty'}), 400
    
    try:
        # Generate change ID
        change_id = str(uuid.uuid4())
        
        # Get original content from issue - handle different data structures
        issue = accessibility_issues[issue_id]
        original_content = ''
        
        if isinstance(issue.get('details'), dict):
            original_content = issue['details'].get('original_content', '') or issue['details'].get('content', '')
        elif isinstance(issue.get('details'), list) and issue['details']:
            original_content = issue['details'][0]
        else:
            original_content = issue.get('description', '')
        
        # Check for duplicate staging of same content
        existing_changes = [
            change for change in staged_changes.values()
            if (change['issue_id'] == issue_id and 
                change['new_content'] == new_content and 
                change['status'] == 'staged')
        ]
        
        if existing_changes:
            return jsonify({
                'error': 'This change is already staged',
                'existing_change_id': existing_changes[0]['id']
            }), 409
        
        # Calculate detailed diff
        diff = calculate_diff(original_content, new_content)
        
        # Create staged change
        staged_changes[change_id] = {
            'id': change_id,
            'issue_id': issue_id,
            'document_id': issue['document_id'],
            'original_content': original_content,
            'new_content': new_content,
            'change_type': data.get('change_type', 'manual'),
            'element_xpath': issue.get('element_xpath', ''),  # Pass XPath from issue
            'created_at': datetime.now().isoformat(),
            'status': 'staged',
            'diff': diff
        }
        
        return jsonify({
            'change_id': change_id,
            'issue_id': issue_id,
            'document_id': issue['document_id'],
            'diff': diff,
            'status': 'staged',
            'created_at': staged_changes[change_id]['created_at']
        }), 201
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/documents/<document_id>/apply-changes', methods=['POST'])
def apply_changes(document_id):
    """Apply all staged changes to document"""
    if document_id not in documents:
        return jsonify({'error': 'Document not found'}), 404
    
    # Optional: Allow applying specific changes via request body
    data = request.get_json() or {}
    specific_change_ids = data.get('change_ids', [])
    
    try:
        # Get all staged changes for this document
        if specific_change_ids:
            document_changes = [
                change for change in staged_changes.values() 
                if (change['document_id'] == document_id and 
                    change['status'] == 'staged' and 
                    change['id'] in specific_change_ids)
            ]
        else:
            document_changes = [
                change for change in staged_changes.values() 
                if change['document_id'] == document_id and change['status'] == 'staged'
            ]
        
        if not document_changes:
            error_msg = 'No staged changes found for this document'
            if specific_change_ids:
                error_msg = 'No matching staged changes found for the specified change IDs'
            return jsonify({'error': error_msg}), 400
        
        # Validate that all changes are still valid
        invalid_changes = []
        for change in document_changes:
            if change['issue_id'] not in accessibility_issues:
                invalid_changes.append({
                    'change_id': change['id'],
                    'reason': 'Associated issue no longer exists'
                })
        
        if invalid_changes:
            return jsonify({
                'error': 'Some changes cannot be applied',
                'invalid_changes': invalid_changes
            }), 400
        
        # Get the document file path
        document = documents[document_id]
        file_path = document['file_path']
        
        # Apply changes to the actual DOCX file
        docx_result = apply_changes_to_docx(file_path, document_changes)
        
        if not docx_result['success']:
            return jsonify({
                'error': 'Failed to apply changes to document',
                'details': docx_result.get('error', 'Unknown error'),
                'failed_changes': docx_result.get('failed_changes', [])
            }), 500
        
        # Generate new document ID for updated version (for tracking)
        updated_document_id = str(uuid.uuid4())
        applied_timestamp = datetime.now().isoformat()
        
        # Update staged changes status based on actual DOCX modification results
        successfully_applied = docx_result['applied_changes']
        failed_to_apply = docx_result['failed_changes']
        
        applied_changes = []
        change_summaries = []
        
        for change in document_changes:
            if change['id'] in successfully_applied:
                change['status'] = 'applied'
                change['applied_at'] = applied_timestamp
                applied_changes.append(change['id'])
                
                # Create summary for response
                change_summaries.append({
                    'change_id': change['id'],
                    'issue_id': change['issue_id'],
                    'change_type': change['change_type'],
                    'diff_summary': change.get('diff', {}).get('summary', {}),
                    'applied_at': applied_timestamp,
                    'status': 'success'
                })
                
                # Mark associated issue as fixed
                if change['issue_id'] in accessibility_issues:
                    accessibility_issues[change['issue_id']]['is_fixed'] = True
                    accessibility_issues[change['issue_id']]['fixed_at'] = applied_timestamp
            else:
                # Change failed to apply
                change['status'] = 'failed'
                change['failed_at'] = applied_timestamp
                
                # Find the failure reason
                failure_reason = 'Unknown error'
                for failed_change in failed_to_apply:
                    if failed_change['change_id'] == change['id']:
                        failure_reason = failed_change['reason']
                        break
                
                change_summaries.append({
                    'change_id': change['id'],
                    'issue_id': change['issue_id'],
                    'change_type': change['change_type'],
                    'status': 'failed',
                    'error': failure_reason,
                    'failed_at': applied_timestamp
                })
        
        # Update document status
        documents[document_id]['status'] = 'remediated'
        documents[document_id]['remediated_at'] = applied_timestamp
        documents[document_id]['applied_changes'] = applied_changes
        
        # Calculate remediation statistics
        total_issues = len([issue for issue in accessibility_issues.values() 
                           if issue['document_id'] == document_id])
        fixed_issues = len([issue for issue in accessibility_issues.values() 
                           if issue['document_id'] == document_id and issue.get('is_fixed', False)])
        
        return jsonify({
            'success': True,
            'updated_document_id': updated_document_id,
            'applied_changes': change_summaries,
            'total_changes': len(applied_changes),
            'docx_modification': {
                'file_modified': True,
                'backup_created': docx_result.get('backup_path') is not None,
                'backup_path': docx_result.get('backup_path'),
                'successfully_applied': len(successfully_applied),
                'failed_to_apply': len(failed_to_apply),
                'modification_details': docx_result
            },
            'remediation_stats': {
                'total_issues': total_issues,
                'fixed_issues': fixed_issues,
                'completion_rate': round((fixed_issues / total_issues * 100), 2) if total_issues > 0 else 0
            },
            'document_status': 'remediated',
            'applied_at': applied_timestamp,
            'message': f'Successfully applied {len(successfully_applied)} changes to the DOCX document. {len(failed_to_apply)} changes failed to apply.' if failed_to_apply else f'Successfully applied all {len(successfully_applied)} changes to the DOCX document.'
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/changes/<change_id>', methods=['DELETE'])
def cancel_staged_change(change_id):
    """Cancel a staged change"""
    if change_id not in staged_changes:
        return jsonify({'error': 'Change not found'}), 404
    
    try:
        change = staged_changes[change_id]
        if change['status'] != 'staged':
            return jsonify({'error': 'Can only cancel staged changes'}), 400
        
        # Remove the change
        del staged_changes[change_id]
        
        return jsonify({
            'change_id': change_id,
            'status': 'cancelled'
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/changes/<document_id>', methods=['GET'])
def get_staged_changes(document_id):
    """Get all staged changes for a document"""
    if document_id not in documents:
        return jsonify({'error': 'Document not found'}), 404
    
    # Filter changes for this document
    document_changes = [
        change for change in staged_changes.values() 
        if change['document_id'] == document_id
    ]
    
    # Separate by status
    staged_changes_list = [c for c in document_changes if c['status'] == 'staged']
    applied_changes_list = [c for c in document_changes if c['status'] == 'applied']
    
    return jsonify({
        'document_id': document_id,
        'staged_changes': staged_changes_list,
        'applied_changes': applied_changes_list,
        'summary': {
            'total_changes': len(document_changes),
            'staged_count': len(staged_changes_list),
            'applied_count': len(applied_changes_list)
        }
    })

@app.route('/api/changes/<document_id>/clear', methods=['POST'])
def clear_staged_changes(document_id):
    """Clear all staged changes for a document"""
    if document_id not in documents:
        return jsonify({'error': 'Document not found'}), 404
    
    try:
        # Find all staged changes for this document
        changes_to_remove = [
            change_id for change_id, change in staged_changes.items()
            if change['document_id'] == document_id and change['status'] == 'staged'
        ]
        
        # Remove the changes
        for change_id in changes_to_remove:
            del staged_changes[change_id]
        
        return jsonify({
            'success': True,
            'cleared_changes': len(changes_to_remove),
            'message': f'Cleared {len(changes_to_remove)} staged changes'
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/changes/<change_id>/preview', methods=['GET'])
def preview_change(change_id):
    """Get detailed preview of a specific change"""
    if change_id not in staged_changes:
        return jsonify({'error': 'Change not found'}), 404
    
    try:
        change = staged_changes[change_id]
        
        # Get associated issue details
        issue = accessibility_issues.get(change['issue_id'], {})
        
        return jsonify({
            'change': change,
            'issue_context': {
                'id': issue.get('id'),
                'clause': issue.get('clause'),
                'description': issue.get('description'),
                'wcag_level': issue.get('wcag_level'),
                'element_xpath': issue.get('element_xpath')
            },
            'diff_details': change.get('diff', {}),
            'can_apply': change['status'] == 'staged'
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/documents/<document_id>/download', methods=['GET'])
def download_modified_document(document_id):
    """Download the modified DOCX document"""
    if document_id not in documents:
        return jsonify({'error': 'Document not found'}), 404
    
    try:
        document = documents[document_id]
        file_path = document['file_path']
        filename = document['filename']
        
        # Add 'modified' to filename if document has been remediated
        if document.get('status') == 'remediated':
            name_parts = filename.rsplit('.', 1)
            modified_filename = f"{name_parts[0]}_modified.{name_parts[1]}"
        else:
            modified_filename = filename
        
        return send_file(
            file_path,
            as_attachment=True,
            download_name=modified_filename,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/documents/<document_id>/restore', methods=['POST'])
def restore_document_backup(document_id):
    """Restore document from backup"""
    if document_id not in documents:
        return jsonify({'error': 'Document not found'}), 404
    
    try:
        document = documents[document_id]
        file_path = document['file_path']
        backup_path = file_path.replace('.docx', '_backup.docx')
        
        if not os.path.exists(backup_path):
            return jsonify({'error': 'No backup found for this document'}), 404
        
        # Restore from backup
        shutil.copy2(backup_path, file_path)
        
        # Reset document status
        documents[document_id]['status'] = 'ready'
        documents[document_id]['restored_at'] = datetime.now().isoformat()
        
        # Reset associated issues
        for issue_id, issue in accessibility_issues.items():
            if issue['document_id'] == document_id:
                issue['is_fixed'] = False
                if 'fixed_at' in issue:
                    del issue['fixed_at']
        
        # Reset staged changes for this document
        for change_id, change in list(staged_changes.items()):
            if change['document_id'] == document_id and change['status'] == 'applied':
                change['status'] = 'reverted'
                change['reverted_at'] = datetime.now().isoformat()
        
        return jsonify({
            'success': True,
            'message': 'Document restored from backup successfully',
            'document_status': 'ready',
            'restored_at': documents[document_id]['restored_at']
        })
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500

if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)